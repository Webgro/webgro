"""
Build a Xero Advanced Invoice Template as a .docx file.

Design system mirrors the Webgro proposal deck:
  - Page background: wg-mist cream (#F4F6FB)
  - Primary ink: wg-ink #07080C for text
  - Triad accent: blue #2D8DFF / violet #7C3AED / teal #00C9A7
  - Display font: Space Grotesk (fallback Arial)
  - Body font: Inter (fallback Arial)
  - Mono labels: JetBrains Mono (fallback Consolas)

Xero Advanced templates use merge fields wrapped in guillemets («FieldName»).
Python-docx doesn't handle Word's native MERGEFIELD instructions well enough
for Xero's parser, so we keep the literal «...» text in place — Xero's
server-side renderer substitutes on invoice generation.

Reference: https://central.xero.com/s/article/Advanced-invoice-template-guide

Upload:
  Xero → Settings → Invoice settings → New Branding Theme → Upload DOCX.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from pathlib import Path

# ---------------------------------------------------------------------------
# Palette (matches globals.css tokens)
# ---------------------------------------------------------------------------
INK = RGBColor(0x07, 0x08, 0x0C)
INK_SOFT = RGBColor(0x52, 0x5B, 0x6E)
SLATE = RGBColor(0x8B, 0x94, 0xA8)
MIST = "F4F6FB"           # hex without # for xml shading
WHITE = "FFFFFF"
LINE = "E6EAF2"           # hairline grey
# RGBColor white for run text colour — used for making the hidden
# TableStart/TableEnd marker merge fields invisible.
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
BLUE = RGBColor(0x2D, 0x8D, 0xFF)
VIOLET = RGBColor(0x7C, 0x3A, 0xED)
TEAL = RGBColor(0x00, 0xC9, 0xA7)

# Fonts chosen to render cleanly inside Xero's PDF engine, which doesn't
# have Space Grotesk / Inter / JetBrains Mono available. We pick fonts
# that ship with the Xero renderer (and Windows / macOS Word) so the
# preview and the rendered PDF look the same, rather than falling back
# silently to Times New Roman (which is what happened on the first try).
#
# Trebuchet MS is the closest "humanist modern sans" equivalent to Space
# Grotesk in a Word-safe font. Calibri is the closest clean body sans to
# Inter. Consolas is a solid mono fallback.
DISPLAY_FONT = "Trebuchet MS"
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

OUT = Path(
    "/Users/michaelbroadbridge/Desktop/Claude/webgro-site/xero/"
    "webgro-invoice-template.docx"
)
LOGO = Path(
    "/Users/michaelbroadbridge/Desktop/Claude/webgro-site/public/brand/logo.png"
)


# ---------------------------------------------------------------------------
# Low-level XML helpers: python-docx doesn't expose page colour, table cell
# shading, single-side borders, or colored paragraph borders cleanly, so we
# drop to lxml for those.
# ---------------------------------------------------------------------------


def set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_vertical_align(cell, value: str = "center") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), value)
    tcPr.append(va)


def remove_all_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    # Replace any existing borders element
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def set_cell_border(cell, edge: str, color: str, size: int = 4) -> None:
    """Single-side cell border. `edge` in top/left/bottom/right."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    b = OxmlElement(f"w:{edge}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))  # eighths of a point
    b.set(qn("w:space"), "0")
    b.set(qn("w:color"), color)
    tcBorders.append(b)


def set_page_background(doc, hex_color: str) -> None:
    """Cream page background. Some renderers (including Xero's PDF engine)
    respect this, others don't. Fallback: we shade the outer table."""
    settings = doc.settings.element
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_color)
    # background must be the first child of document for Word to render it
    root = doc.element
    # prepend
    root.insert(0, bg)
    # Also enable display via settings
    display = OxmlElement("w:displayBackgroundShape")
    settings.append(display)


# ---------------------------------------------------------------------------
# Text-level helpers
# ---------------------------------------------------------------------------


def add_run(paragraph, text, *, font=BODY_FONT, size=10, bold=False,
            italic=False, color: RGBColor = INK, tracking_pts: float | None = None):
    run = paragraph.add_run(text)
    run.font.name = font
    # East-Asia font fallback so Word doesn't override
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if tracking_pts is not None:
        # spacing in 20ths of a point
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(int(tracking_pts * 20)))
        rPr.append(spacing)
    return run


def add_merge(paragraph, field_name: str, *, font=BODY_FONT, size=10,
              bold=False, italic=False, color: RGBColor = INK,
              tracking_pts: float | None = None,
              format_switch: str | None = None):
    """Add a proper Word MERGEFIELD that Xero's parser recognises.

    `format_switch` is an optional Word merge-field format string, e.g.
    `\\# "£#,##0.00"` for currency, `\\@ "dd MMM yyyy"` for dates.
    Xero honours these at render time — necessary if you want amounts
    prefixed with currency symbols or dates in a human format rather
    than the raw ISO date Xero stores internally.
    """
    display = f"«{field_name}»"
    run = add_run(paragraph, display, font=font, size=size, bold=bold,
                  italic=italic, color=color, tracking_pts=tracking_pts)
    r_elem = run._r

    parent = r_elem.getparent()
    idx = list(parent).index(r_elem)
    parent.remove(r_elem)

    fldSimple = OxmlElement("w:fldSimple")
    instr = f' MERGEFIELD  {field_name}  '
    if format_switch:
        instr += f'{format_switch}  '
    instr += '\\* MERGEFORMAT '
    fldSimple.set(qn("w:instr"), instr)
    fldSimple.append(r_elem)
    parent.insert(idx, fldSimple)
    return run


def set_paragraph_spacing(p, *, before=0, after=4, line_rule_auto=True):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_rule_auto:
        pf.line_spacing = 1.2


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------


def main() -> None:
    doc = Document()

    # ------ page setup -----------------------------------------------------
    # A4, trimmed margins to claw back a centimetre each side for line items.
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    set_page_background(doc, MIST)

    # ------ HEADER BLOCK ---------------------------------------------------
    # 2-col table: left = Webgro wordmark + tagline, right = big "INVOICE"
    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_all_table_borders(header)

    # column widths
    header.columns[0].width = Cm(10.0)
    header.columns[1].width = Cm(7.4)
    for row in header.rows:
        row.cells[0].width = Cm(10.0)
        row.cells[1].width = Cm(7.4)

    # LEFT cell — real logo image + tagline
    left = header.rows[0].cells[0]
    set_cell_shading(left, MIST)
    p = left.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Add the logo as a picture run. Width picked so the wordmark sits at
    # roughly the same visual weight as the old 40pt "Webgro" text.
    logo_run = p.add_run()
    if LOGO.exists():
        logo_run.add_picture(str(LOGO), width=Cm(4.5))
    else:
        add_run(p, "Webgro", font=DISPLAY_FONT, size=36, bold=True, color=INK)

    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(0)
    add_run(
        p2,
        "AI-FIRST eCOMMERCE & WORDPRESS AGENCY",
        font=MONO_FONT,
        size=7,
        color=SLATE,
        tracking_pts=1.2,
    )

    # RIGHT cell — big INVOICE label + number
    right = header.rows[0].cells[1]
    set_cell_shading(right, MIST)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "[ INVOICE ]", font=MONO_FONT, size=9, color=BLUE,
            tracking_pts=2.0)

    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    add_merge(p2, "InvoiceNumber", font=DISPLAY_FONT, size=24, bold=True,
              color=INK)

    # ------ GRADIENT ACCENT BAR -------------------------------------------
    # Tri-column table, each cell a thin colored strip.
    bar = doc.add_table(rows=1, cols=3)
    bar.autofit = False
    remove_all_table_borders(bar)
    bar_cells = bar.rows[0].cells
    bar_colors = ["2D8DFF", "7C3AED", "00C9A7"]
    for cell, col in zip(bar_cells, bar_colors):
        cell.width = Cm(5.8)
        set_cell_shading(cell, col)
        # empty paragraph but force minimal height
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("")
        r.font.size = Pt(2)
    # tiny spacer
    _sp = doc.add_paragraph()
    _sp.paragraph_format.space_before = Pt(0)
    _sp.paragraph_format.space_after = Pt(6)

    # ------ FROM / BILL TO -------------------------------------------------
    ft = doc.add_table(rows=1, cols=2)
    ft.autofit = False
    remove_all_table_borders(ft)
    for row in ft.rows:
        row.cells[0].width = Cm(8.5)
        row.cells[1].width = Cm(8.9)

    # FROM
    fcell = ft.rows[0].cells[0]
    set_cell_shading(fcell, WHITE)
    set_cell_border(fcell, "left", "2D8DFF", size=16)
    p = fcell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    # left padding via indent (python-docx has no cell padding for single side)
    p.paragraph_format.left_indent = Cm(0.35)
    add_run(p, "FROM", font=MONO_FONT, size=8, color=BLUE, tracking_pts=1.6)

    # Webgro's own details are static (our own org info doesn't change per
    # invoice), EXCEPT the VAT number which is pulled from Xero's
    # organisation settings via the OrganisationTaxNumber merge field —
    # change it in one place in Xero instead of having to rebuild the
    # template when we register for VAT / change the number.
    lines = [
        ("Webgro Ltd", DISPLAY_FONT, 12, True, INK, False),
        ("12 Longshot Lane", BODY_FONT, 9, False, INK_SOFT, False),
        ("Bracknell, Berkshire, RG12 1RL", BODY_FONT, 9, False, INK_SOFT, False),
        ("United Kingdom", BODY_FONT, 9, False, INK_SOFT, False),
        ("", BODY_FONT, 4, False, INK, False),
        ("Company No. 10889889", MONO_FONT, 7, False, SLATE, False),
    ]
    for text, font, size, bold, color, _ in lines:
        para = fcell.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.35)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(1)
        add_run(para, text, font=font, size=size, bold=bold, color=color)

    # VAT line with merge field for the number
    vat = fcell.add_paragraph()
    vat.paragraph_format.left_indent = Cm(0.35)
    vat.paragraph_format.space_before = Pt(0)
    vat.paragraph_format.space_after = Pt(1)
    add_run(vat, "VAT No. ", font=MONO_FONT, size=7, color=SLATE)
    add_merge(vat, "OrganisationTaxNumber", font=MONO_FONT, size=7, color=SLATE)

    spacer = fcell.add_paragraph()
    spacer.paragraph_format.left_indent = Cm(0.35)
    spacer.paragraph_format.space_after = Pt(1)
    add_run(spacer, "", font=BODY_FONT, size=4, color=INK)

    # Contact
    for text, font, size, bold, color in [
        ("hello@webgro.co.uk", BODY_FONT, 9, False, INK),
        ("+44 (0) 1344 231 119", BODY_FONT, 9, False, INK),
    ]:
        para = fcell.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.35)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(1)
        add_run(para, text, font=font, size=size, bold=bold, color=color)

    # small bottom padding
    para = fcell.add_paragraph()
    para.paragraph_format.space_after = Pt(8)

    # BILL TO
    tcell = ft.rows[0].cells[1]
    set_cell_shading(tcell, WHITE)
    set_cell_border(tcell, "left", "7C3AED", size=16)
    p = tcell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "BILL TO", font=MONO_FONT, size=8, color=VIOLET,
            tracking_pts=1.6)

    # Contact name (the client's display name on the invoice)
    name_p = tcell.add_paragraph()
    name_p.paragraph_format.left_indent = Cm(0.35)
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(1)
    add_merge(name_p, "ContactName", font=DISPLAY_FONT, size=12, bold=True,
              color=INK)

    # Full postal address — Xero populates this as a multi-line string,
    # so one merge field on its own paragraph renders the whole address
    # block with line breaks preserved.
    addr_p = tcell.add_paragraph()
    addr_p.paragraph_format.left_indent = Cm(0.35)
    addr_p.paragraph_format.space_before = Pt(0)
    addr_p.paragraph_format.space_after = Pt(1)
    add_merge(addr_p, "ContactPostalAddress", font=BODY_FONT, size=9,
              color=INK_SOFT)

    # Spacer then email
    sep = tcell.add_paragraph()
    sep.paragraph_format.left_indent = Cm(0.35)
    sep.paragraph_format.space_after = Pt(1)
    add_run(sep, "", font=BODY_FONT, size=4, color=INK)

    email_p = tcell.add_paragraph()
    email_p.paragraph_format.left_indent = Cm(0.35)
    email_p.paragraph_format.space_before = Pt(0)
    email_p.paragraph_format.space_after = Pt(1)
    add_merge(email_p, "ContactEmailAddress", font=BODY_FONT, size=9,
              color=INK)

    pad = tcell.add_paragraph()
    pad.paragraph_format.space_after = Pt(8)

    # tighter spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)

    # ------ METADATA STRIP: Issue Date / Due Date / Reference -------------
    meta = doc.add_table(rows=2, cols=3)
    meta.autofit = False
    remove_all_table_borders(meta)
    meta_colors = [BLUE, VIOLET, TEAL]
    meta_labels = ["ISSUED", "DUE", "REFERENCE"]
    # (field_name, format_switch) — dates use the \@ date-format switch
    # so they render as "21 Apr 2026" rather than ISO 2026-04-21.
    DATE_FMT = r'\@ "dd MMM yyyy"'
    meta_specs = [
        ("InvoiceDate", DATE_FMT),
        ("InvoiceDueDate", DATE_FMT),
        ("Reference", None),
    ]
    for col in range(3):
        hdr = meta.rows[0].cells[col]
        val = meta.rows[1].cells[col]
        hdr.width = Cm(5.8)
        val.width = Cm(5.8)
        set_cell_shading(hdr, WHITE)
        set_cell_shading(val, WHITE)
        set_cell_border(hdr, "top", "E6EAF2", size=6)
        set_cell_border(val, "bottom", "E6EAF2", size=6)
        p = hdr.paragraphs[0]
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, meta_labels[col], font=MONO_FONT, size=7,
                color=meta_colors[col], tracking_pts=1.6)
        p2 = val.paragraphs[0]
        p2.paragraph_format.left_indent = Cm(0.35)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        field_name, fmt = meta_specs[col]
        add_merge(p2, field_name, font=DISPLAY_FONT, size=12,
                  bold=True, color=INK, format_switch=fmt)

    # tighter spacer before line items
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

    # ------ LINE ITEMS TABLE ----------------------------------------------
    # Xero auto-replicates the body row based on structure, so we create
    # header + ONE body row (marked so it repeats).
    items = doc.add_table(rows=2, cols=4)
    items.autofit = False
    remove_all_table_borders(items)
    widths = [Cm(9.5), Cm(2.1), Cm(2.8), Cm(3.0)]
    for i, w in enumerate(widths):
        items.columns[i].width = w
        for row in items.rows:
            row.cells[i].width = w

    # --- header row ---
    h = items.rows[0]
    h_titles = ["DESCRIPTION", "QTY", "UNIT", "AMOUNT"]
    h_aligns = [
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT,
    ]
    for idx, (cell, title, align) in enumerate(zip(h.cells, h_titles, h_aligns)):
        set_cell_shading(cell, WHITE)
        set_cell_border(cell, "bottom", "07080C", size=10)
        set_cell_vertical_align(cell, "center")
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Cm(0.25) if align == WD_ALIGN_PARAGRAPH.LEFT else Cm(0)
        p.paragraph_format.right_indent = Cm(0.25) if align == WD_ALIGN_PARAGRAPH.RIGHT else Cm(0)
        add_run(p, title, font=MONO_FONT, size=8, color=INK, bold=True,
                tracking_pts=1.4)

    # --- body row (Xero replicates this row per line item) ---
    #
    # CRITICAL: both TableStart:LineItem AND TableEnd:LineItem MUST live
    # in THIS same row. If TableStart ends up in the header row above,
    # Xero can't pair them and PDF generation fails. (That's the bug the
    # previous version hit.) We stash them in cell 1 and cell N at 1pt
    # white-on-white so they're present in the XML but invisible.
    b = items.rows[1]

    # (field_name, alignment, font, size, bold, color, format_switch)
    body_fields = [
        ("Description", WD_ALIGN_PARAGRAPH.LEFT, BODY_FONT, 10, False, INK, None),
        ("Quantity", WD_ALIGN_PARAGRAPH.RIGHT, MONO_FONT, 10, False, INK_SOFT,
         r'\# #,##0.##'),
        ("UnitAmount", WD_ALIGN_PARAGRAPH.RIGHT, MONO_FONT, 10, False, INK_SOFT,
         r'\# "#,##0.00;(#,##0.00)"'),
        ("LineAmount", WD_ALIGN_PARAGRAPH.RIGHT, DISPLAY_FONT, 11, True, INK,
         r'\# "#,##0.00;(#,##0.00)"'),
    ]
    for idx, (cell, (field, align, font, size, bold, color, fmt)) in enumerate(
        zip(b.cells, body_fields)
    ):
        set_cell_shading(cell, WHITE)
        set_cell_border(cell, "bottom", "E6EAF2", size=4)
        set_cell_vertical_align(cell, "center")
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.25) if align == WD_ALIGN_PARAGRAPH.LEFT else Cm(0)
        p.paragraph_format.right_indent = Cm(0.25) if align == WD_ALIGN_PARAGRAPH.RIGHT else Cm(0)

        if idx == 0:
            # TableStart goes INSIDE the body row before the first
            # visible line-item field.
            add_merge(p, "TableStart:LineItem", font=MONO_FONT, size=1,
                      color=WHITE_RGB)

        add_merge(p, field, font=font, size=size, bold=bold, color=color,
                  format_switch=fmt)

        if idx == len(body_fields) - 1:
            # TableEnd comes after the last visible line-item field in
            # the same row.
            add_merge(p, "TableEnd:LineItem", font=MONO_FONT, size=1,
                      color=WHITE_RGB)

    # small spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

    # ------ FOOTER BLOCK: Pay + Thanks (LEFT) | Totals (RIGHT) ------------
    # Single row, two columns. Left cell stacks bank details + thanks copy.
    # Right cell carries the subtotals and the highlighted Amount Due.
    footer = doc.add_table(rows=1, cols=2)
    footer.autofit = False
    remove_all_table_borders(footer)
    footer.columns[0].width = Cm(10.5)
    footer.columns[1].width = Cm(7.5)
    for row in footer.rows:
        row.cells[0].width = Cm(10.5)
        row.cells[1].width = Cm(7.5)

    # ---------- LEFT: payment + thanks (stacked) ----------
    pl = footer.rows[0].cells[0]
    set_cell_shading(pl, WHITE)
    set_cell_border(pl, "left", "00C9A7", size=16)

    p = pl.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "HOW TO PAY", font=MONO_FONT, size=8, color=TEAL,
            tracking_pts=1.6)

    # Static payment detail lines (hardcoded, no merge fields)
    static_pay = [
        ("Bank transfer · Monzo", DISPLAY_FONT, 11, True, INK),
        ("Webgro Ltd", BODY_FONT, 9, False, INK),
        ("Sort code · 04-00-04", BODY_FONT, 9, False, INK_SOFT),
        ("Account · 65693471", BODY_FONT, 9, False, INK_SOFT),
    ]
    for text, font, size, bold, color in static_pay:
        para = pl.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.35)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(1)
        add_run(para, text, font=font, size=size, bold=bold, color=color)

    # Reference line — label is static, the invoice number is a real
    # merge field so it populates on send.
    pref = pl.add_paragraph()
    pref.paragraph_format.left_indent = Cm(0.35)
    pref.paragraph_format.space_before = Pt(0)
    pref.paragraph_format.space_after = Pt(1)
    add_run(pref, "Reference · ", font=BODY_FONT, size=9, color=INK_SOFT)
    add_merge(pref, "InvoiceNumber", font=BODY_FONT, size=9, color=INK_SOFT)

    # Blank spacer line (keeps the visual rhythm between "Reference" and
    # the divider before Thanks). Payment terms already appear in the
    # "DUE" date of the metadata strip above, so we don't repeat them
    # here.
    sp_line = pl.add_paragraph()
    sp_line.paragraph_format.left_indent = Cm(0.35)
    sp_line.paragraph_format.space_after = Pt(1)
    add_run(sp_line, "", font=BODY_FONT, size=3, color=INK)

    # Divider paragraph inside the cell
    dp = pl.add_paragraph()
    dp.paragraph_format.left_indent = Cm(0.35)
    dp.paragraph_format.right_indent = Cm(0.35)
    dp.paragraph_format.space_before = Pt(6)
    dp.paragraph_format.space_after = Pt(4)
    pPr = dp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E6EAF2")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Thanks + short note
    pt = pl.add_paragraph()
    pt.paragraph_format.left_indent = Cm(0.35)
    pt.paragraph_format.space_before = Pt(2)
    pt.paragraph_format.space_after = Pt(2)
    add_run(pt, "Thanks.", font=DISPLAY_FONT, size=22, bold=True,
            italic=True, color=INK)

    pn = pl.add_paragraph()
    pn.paragraph_format.left_indent = Cm(0.35)
    pn.paragraph_format.right_indent = Cm(0.35)
    pn.paragraph_format.space_before = Pt(2)
    pn.paragraph_format.space_after = Pt(2)
    add_run(pn,
            "Questions or changes? Email hello@webgro.co.uk and we'll sort it.",
            font=BODY_FONT, size=8, color=INK_SOFT)

    pm = pl.add_paragraph()
    pm.paragraph_format.left_indent = Cm(0.35)
    pm.paragraph_format.right_indent = Cm(0.35)
    pm.paragraph_format.space_before = Pt(6)
    pm.paragraph_format.space_after = Pt(8)
    add_run(pm, "WEBGRO LTD  ·  REG 10889889  ·  PART OF BROADBRIDGE GROUP",
            font=MONO_FONT, size=7, color=SLATE, tracking_pts=1.2)

    # ---------- RIGHT: totals ----------
    tr = footer.rows[0].cells[1]
    set_cell_shading(tr, WHITE)
    set_cell_vertical_align(tr, "top")

    # Xero's real field names for invoice totals. `AmountPaid` isn't
    # exposed as a template field, so we omit a Paid row; InvoiceAmountDue
    # is the outstanding balance after any part-payments.
    MONEY_FMT = r'\# "#,##0.00;(#,##0.00)"'
    totals_rows = [
        ("Subtotal", "InvoiceSubTotal"),
        ("VAT", "TaxTotal"),
    ]

    first = True
    for label, field in totals_rows:
        p = tr.paragraphs[0] if first else tr.add_paragraph()
        is_first = first
        first = False
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.right_indent = Cm(0.35)
        p.paragraph_format.space_before = Pt(8 if is_first else 0)
        p.paragraph_format.space_after = Pt(3)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(6.7), WD_ALIGN_PARAGRAPH.RIGHT)
        add_run(p, label, font=BODY_FONT, size=9, color=INK_SOFT)
        add_run(p, "\t", font=BODY_FONT, size=9)
        add_merge(p, field, font=MONO_FONT, size=10, color=INK,
                  format_switch=MONEY_FMT)

    # Divider
    dv = tr.add_paragraph()
    dv.paragraph_format.left_indent = Cm(0.35)
    dv.paragraph_format.right_indent = Cm(0.35)
    dv.paragraph_format.space_before = Pt(4)
    dv.paragraph_format.space_after = Pt(4)
    pPr = dv._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "07080C")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # AMOUNT DUE row
    pd = tr.add_paragraph()
    pd.paragraph_format.left_indent = Cm(0.35)
    pd.paragraph_format.right_indent = Cm(0.35)
    pd.paragraph_format.space_before = Pt(2)
    pd.paragraph_format.space_after = Pt(2)
    add_run(pd, "AMOUNT DUE", font=MONO_FONT, size=8, color=TEAL, bold=True,
            tracking_pts=1.6)

    pda = tr.add_paragraph()
    pda.paragraph_format.left_indent = Cm(0.35)
    pda.paragraph_format.right_indent = Cm(0.35)
    pda.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pda.paragraph_format.space_before = Pt(0)
    pda.paragraph_format.space_after = Pt(8)
    add_merge(pda, "InvoiceAmountDue", font=DISPLAY_FONT, size=20,
              bold=True, color=INK, format_switch=MONEY_FMT)

    # Save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
